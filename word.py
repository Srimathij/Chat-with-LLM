
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param text: The display text for the hyperlink.
    :param url: A string containing the required url.
    :return: None
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Add a w:rStyle element and set its value to Hyperlink
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    # Create a w:t element and set the text value
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    # Append the w:r element to the hyperlink element
    hyperlink.append(new_run)

    # Append the hyperlink element to the paragraph
    paragraph._p.append(hyperlink)


def create_proposal_doc(quantity, tile_name, size, thickness, price_per_tile):
    try:
        # Create a new Document
        doc = Document()

        # Define a table with 1 row and 2 columns for header (logo on the right)
        header_table = doc.add_table(rows=1, cols=2)
        # Left cell remains empty
        header_table.cell(0, 0).text = ""
        # Right cell: add image (ensure the image file exists in the working directory)
        image_paragraph = header_table.cell(0, 1).paragraphs[0]
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = image_paragraph.add_run()
        run.add_picture("boo.png", width=Inches(2))  # Update image file path if needed

        # Title
        doc.add_heading('Tiles Proposal', level=1)

        # Sender and Recipient Details
        doc.add_paragraph('From: Niro Ceramics')
        doc.add_paragraph('To: Vijai')
        doc.add_paragraph('\n')

        # Niro Overview
        doc.add_paragraph(
            'Pioneer in the tiles industry producing homogeneous tiles, We offer a variety of products, from porcelain and ceramic tiles '
            'to glass mosaics and even bathroom sanitaryware through its various brands.'
        )

        # Proposal Overview
        doc.add_heading('Proposal Overview', level=2)
        doc.add_paragraph(
            'We are delighted to present our proposal for your tile requirements. Below are the details of the order and product specifications:'
        )

        # Product Specifications
        doc.add_heading('Product Specifications', level=2)
        spec_table = doc.add_table(rows=2, cols=2)
        spec_table.style = 'Light Grid'
        spec_table.cell(0, 0).text = 'Size'
        spec_table.cell(0, 1).text = size
        spec_table.cell(1, 0).text = 'Tile Thickness'
        spec_table.cell(1, 1).text = thickness

        # Order Details Calculation
        doc.add_heading('Order Details', level=2)
        quantity = int(quantity)
        price_per_tile = float(price_per_tile)
        gst_percentage = 8  # GST percentage
        total_without_gst = quantity * price_per_tile
        gst_amount = total_without_gst * (gst_percentage / 100)
        total_with_gst = total_without_gst + gst_amount

        # Create invoice-like table for order details
        order_table = doc.add_table(rows=5, cols=4)
        order_table.style = 'Light Grid'
        # Table headers
        order_table.cell(0, 0).text = 'Description'
        order_table.cell(0, 1).text = 'Quantity'
        order_table.cell(0, 2).text = 'Price per Unit'
        order_table.cell(0, 3).text = 'Amount'

        # Product details row
        order_table.cell(1, 0).text = tile_name
        order_table.cell(1, 1).text = str(quantity)
        order_table.cell(1, 2).text = f'RM{price_per_tile:.2f}'
        order_table.cell(1, 3).text = f'RM{total_without_gst:.2f}'

        # GST details row
        order_table.cell(2, 0).text = 'SST (8%)'
        order_table.cell(2, 1).text = ""
        order_table.cell(2, 2).text = f'{gst_percentage}%'
        order_table.cell(2, 3).text = f'RM{gst_amount:.2f}'

        # Total amount row
        order_table.cell(3, 0).text = 'Total Amount (Incl. SST)'
        order_table.cell(3, 1).text = ""
        order_table.cell(3, 2).text = ""
        order_table.cell(3, 3).text = f'RM{total_with_gst:.2f}'

        # Add extra row if needed for spacing (or any additional information)
        order_table.cell(4, 0).text = ""
        order_table.cell(4, 1).text = ""
        order_table.cell(4, 2).text = ""
        order_table.cell(4, 3).text = ""

        # NIRO Terms & Conditions
        doc.add_heading('NIRO Terms & Conditions of Sale', level=2)
        terms_paragraph = doc.add_paragraph()
        terms_paragraph.add_run(
            "This quotation is provided subject to Niro Ceramic Sales & Services (M) Sdn Bhd’s standard terms and conditions of sale which can be found "
        )
        # Add the hyperlink for terms and conditions
        add_hyperlink(terms_paragraph, "here", "https://niroceramic.com/media/general-purchasing-terms-and-conditions/")
        terms_paragraph.add_run(
            ". These terms and conditions apply to this quotation and any subsequent order(s) notwithstanding anything to the contrary contained in or incorporated into any document from or oral statement made by you, the customer. No variation or amendment to the conditions shall be of any effect unless expressly agreed, in writing, by a person authorised to sign on behalf of Niro Ceramic Sales & Services (M) Sdn Bhd. By accepting this quotation, I confirm that I have read and I unconditionally accept the terms and conditions of Niro Ceramic Sales & Services (M) Sdn Bhd and I am authorised to enter into this contract."
        )
        terms_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph('\n')

        # Contact Information
        doc.add_heading('Contact Us', level=2)
        doc.add_paragraph('For further inquiries or assistance, feel free to reach out to us at:')

        contact_paragraph = doc.add_paragraph()
        run = contact_paragraph.add_run('- Email: ')
        run.bold = True
        contact_paragraph.add_run('nirogranite@nirogroup.com')

        contact_paragraph = doc.add_paragraph()
        run = contact_paragraph.add_run('- Address: ')
        run.bold = True
        contact_paragraph.add_run('Niro Ceramic Group Headquarters Lot 2, Persiaran Sultan, Seksyen 15, 40200 Shah Alam, Selangor, Malaysia.')

        contact_paragraph = doc.add_paragraph()
        run = contact_paragraph.add_run('- Phone: ')
        run.bold = True
        contact_paragraph.add_run('+603 5033 9333')

        # Closing Note
        doc.add_paragraph('\n')
        doc.add_paragraph('Thank you for considering Niro Ceramics. We look forward to a successful collaboration!')
        closing_paragraph = doc.add_paragraph()
        run = closing_paragraph.add_run('Warm Regards,')
        run.bold = True
        doc.add_paragraph('Niro Ceramics Team')
        doc.add_paragraph('Bringing joy to people through their living spaces.')

        # Save the Document
        doc.save('Tiles_Proposal.docx')
        print("Document saved successfully as 'Tiles_Proposal.docx'.")

    except Exception as e:
        print(f"An error occurred: {e}")


# Example usage:
if __name__ == "__main__":
    # Sample input values
    quantity = 100
    tile_name = "Premium Porcelain Tile"
    size = "600x600 mm"
    thickness = "10 mm"
    price_per_tile = 15.50

    create_proposal_doc(quantity, tile_name, size, thickness, price_per_tile)
